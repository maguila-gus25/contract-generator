import os
import unicodedata
from docx import Document
from fpdf import FPDF
from ..models import Contract

def remover_acentos(texto: str) -> str:
    return unicodedata.normalize('NFKD', texto).encode('ASCII', 'ignore').decode('ASCII')

def format_qualificacao(data: dict) -> str:
    if not data:
        return "[DADOS NÃO INFORMADOS]"
    
    nome = data.get("name", "[NOME NÃO INFORMADO]")
    doc = data.get("document", "[DOCUMENTO NÃO INFORMADO]")
    endereco = data.get("address", "[ENDEREÇO NÃO INFORMADO]")
    numero = data.get("number", "")
    cidade = data.get("city_state", "")
    
    end_completo = f"{endereco}, {numero} - {cidade}" if numero else endereco
    
    return f"{nome}, inscrito(a) sob o documento {doc}, residente e domiciliado(a) em {end_completo}"

def get_contract_text(contract: Contract) -> str:
    c_info = contract.contract_info or {}
    
    texto = f"{contract.title.upper() if contract.title else 'CONTRATO'}\n\n"
    texto += "Por este instrumento particular, de um lado:\n"
    texto += f"{format_qualificacao(contract.contractor_data)}.\n\n"
    texto += "E de outro lado:\n"
    texto += f"{format_qualificacao(contract.contractee_data)}.\n\n"
    
    texto += "Firmam o presente contrato, mediante as cláusulas e condições a seguir:\n\n"
    
    clausulas = c_info.get("clauses", [])
    if not clausulas:
        texto += "CLÁUSULA PRIMEIRA - DO OBJETO\n"
        texto += "As partes concordam com os termos estabelecidos neste instrumento.\n\n"
    else:
        for idx, clausula in enumerate(clausulas, 1):
            titulo = clausula.get("title", f"CLÁUSULA {idx}")
            conteudo = clausula.get("content", "")
            texto += f"{titulo.upper()}\n{conteudo}\n\n"
            
    valor = c_info.get("value", 0.0)
    moeda = c_info.get("currency", "R$")
    pagamento = c_info.get("payment_method", "à vista")
    
    texto += "DO VALOR E PAGAMENTO\n"
    texto += f"O valor total do contrato é de {moeda} {valor:.2f}, pago via {pagamento}.\n\n"
    
    texto += "E por estarem justos e contratados, assinam o presente contrato.\n\n"
    
    data_criacao = contract.created_at.strftime("%d/%m/%Y") if contract.created_at else ""
    texto += f"Data de emissão: {data_criacao}\n\n"
    
    nome_contratante = (contract.contractor_data or {}).get("name", "Contratante")
    nome_contratado = (contract.contractee_data or {}).get("name", "Contratado")
    
    texto += "__________________________________________________\n"
    texto += f"{nome_contratante}\n\n"
    texto += "__________________________________________________\n"
    texto += f"{nome_contratado}\n"
    
    return texto

def generate_pdf(contract: Contract, output_path: str):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # Use Helvetica which is built-in
    pdf.set_font("Helvetica", style="B", size=16)
    
    texto_completo = get_contract_text(contract)
    
    for linha in texto_completo.split('\n'):
        if not linha.strip():
            pdf.ln(5)
            continue
            
        if "DO VALOR" in linha or "Firmam o presente" in linha or "Por este instrumento" in linha:
            pdf.set_font("Helvetica", style="B", size=12)
        elif linha.startswith("CLÁUSULA"):
            pdf.set_font("Helvetica", style="B", size=12)
        else:
            pdf.set_font("Helvetica", size=12)
        
        linha_limpa = remover_acentos(linha)
        try:
            pdf.write(h=8, text=linha_limpa)
            pdf.ln(8)
        except Exception as e:
            print(f"ERRO FPDF NA LINHA: {repr(linha_limpa)}")
            raise
            
    pdf.output(output_path)
    return output_path

def generate_docx(contract: Contract, output_path: str):
    doc = Document()
    texto_completo = get_contract_text(contract)
    
    for linha in texto_completo.split('\n'):
        if not linha.strip():
            continue
            
        p = doc.add_paragraph()
        
        if linha == (contract.title.upper() if contract.title else 'CONTRATO'):
            p.add_run(linha).bold = True
            p.alignment = 1 # Center
        elif linha.startswith("CLÁUSULA") or "DO VALOR" in linha:
            p.add_run(linha).bold = True
        else:
            p.add_run(linha)
            
    doc.save(output_path)
    return output_path
