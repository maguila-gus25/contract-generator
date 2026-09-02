"""Geração do DOCX do contrato.

Acompanha o PDF na **estrutura** — sem número, qualificação em prosa,
cláusulas em romano, itens em dois níveis, fecho com local, data e
assinaturas — mas não no refino tipográfico: o DOCX existe para ser editado
à mão, e um arquivo cheio de formatação manual atrapalha essa edição.
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from ..models.formatters import data_por_extenso
from .base import GeradorBase, rotulo_tipo

# Recuo de um nível de hierarquia, em centímetros.
RECUO = 0.8

PREAMBULO = (
    "As partes acima identificadas têm, entre si, justo e acertado o presente "
    "Contrato, que se regerá pelas cláusulas e condições a seguir descritas."
)


class DocxGenerator(GeradorBase):
    def gerar(self, contrato, caminho_saida: str) -> str:
        caminho = self._preparar_caminho(caminho_saida, "docx")
        doc = Document()
        self._aplicar_estilos(doc)

        dados = contrato.to_dict()
        self._adicionar_cabecalho(doc, contrato)
        self._adicionar_qualificacao(doc, contrato)
        self._adicionar_clausulas(doc, contrato, dados)
        self._adicionar_fecho(doc, contrato, dados)

        doc.save(caminho)
        return caminho

    def _aplicar_estilos(self, doc: Document) -> None:
        estilo = doc.styles["Normal"]
        estilo.font.name = "Arial"
        estilo.font.size = Pt(11)
        paragrafo = estilo.paragraph_format
        paragrafo.space_after = Pt(6)
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _paragrafo(self, doc: Document, texto: str, recuo: float = 0):
        p = doc.add_paragraph(texto)
        if recuo:
            p.paragraph_format.left_indent = Cm(recuo)
        return p

    def _adicionar_cabecalho(self, doc: Document, contrato) -> None:
        titulo = doc.add_heading(f"CONTRATO DE {rotulo_tipo(contrato.tipo)}",
                                 level=1)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _adicionar_qualificacao(self, doc: Document, contrato) -> None:
        for papel, parte in (("CONTRATANTE", contrato.contratante),
                             ("CONTRATADO", contrato.contratado)):
            p = doc.add_paragraph()
            p.add_run(f"{papel}: ").bold = True
            p.add_run(parte.qualificacao())

        doc.add_paragraph(PREAMBULO)

    def _adicionar_clausulas(self, doc: Document, contrato, dados: dict) -> None:
        for clausula in contrato.clausulas:
            titulo = doc.add_paragraph()
            run = titulo.add_run(clausula.cabecalho())
            run.bold = True

            self._paragrafo(doc, clausula.caput(dados))

            for rotulo, texto, subitens in clausula.itens_formatados(dados):
                self._paragrafo(doc, f"{rotulo} {texto}", recuo=RECUO)
                for sub_rotulo, sub_texto in subitens:
                    self._paragrafo(doc, f"{sub_rotulo} {sub_texto}",
                                    recuo=RECUO * 2)

            for texto in clausula.desdobramentos_formatados(dados):
                self._paragrafo(doc, texto, recuo=RECUO)

    def _adicionar_fecho(self, doc: Document, contrato, dados: dict) -> None:
        cidade = dados["contratado_cidade"] or dados["contratante_cidade"]
        local_e_data = data_por_extenso(contrato.data_criacao)
        if cidade:
            local_e_data = f"{cidade}, {local_e_data}"

        doc.add_paragraph()
        p = doc.add_paragraph(local_e_data)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.add_paragraph()

        # Tabela sem bordas: só organiza as duas colunas de assinatura, com a
        # linha desenhada por underscores para o documento continuar editável.
        tabela = doc.add_table(rows=4, cols=2)
        assinaturas = (
            ("CONTRATANTE", dados["contratante_nome"],
             f"{dados['contratante_tipo_documento']} "
             f"{dados['contratante_documento']}"),
            ("CONTRATADO", dados["contratado_nome"],
             f"{dados['contratado_tipo_documento']} "
             f"{dados['contratado_documento']}"),
        )
        for coluna, (papel, nome, documento) in enumerate(assinaturas):
            for linha, texto in enumerate(("_" * 34, nome, documento, papel)):
                tabela.cell(linha, coluna).text = texto

        for row in tabela.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
