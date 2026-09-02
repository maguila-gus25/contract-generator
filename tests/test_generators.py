"""Testes da geração dos documentos: estrutura, numeração e casos de borda."""

import io
import re
from datetime import date

import docx
import pytest

from contract_generator.generators.docx_generator import DocxGenerator
from contract_generator.generators.pdf_generator import PdfGenerator
from contract_generator.models.clause import (Clausula, Item, letra, romano,
                                              carregar_clausulas_padrao)
from contract_generator.models.contract import Contrato
from contract_generator.models.formatters import data_por_extenso
from contract_generator.models.party import Endereco, Parte

ENDERECO_COMPLETO = Endereco("88063-000", "rua Servidão Piloto", "280", "",
                             "Campeche", "Florianópolis", "SC")


def _contrato(contratante, contratado=None, valor=900.0):
    return Contrato(
        tipo="fotografia",
        data_criacao=date(2026, 8, 21),
        data_inicio=date(2026, 8, 21),
        data_fim=None,
        contratante=contratante,
        contratado=contratado or contratante,
        clausulas=carregar_clausulas_padrao("fotografia"),
        valor=valor,
        forma_pagamento="PIX",
        descricao_servico="Ensaio Fotográfico",
        extras={"duracao_horas": "2"},
    )


def _cliente_completo():
    return Parte("Isabel Terra Faillace Holanda", "18133602785", "CPF",
                 "isabelterra09@gmail.com", "(48) 99656-5530",
                 ENDERECO_COMPLETO)


def _cliente_minimo():
    return Parte("Só Nome", "11122233396", "CPF", "", "")


def _texto_do_docx(caminho):
    documento = docx.Document(caminho)
    linhas = [p.text for p in documento.paragraphs]
    for tabela in documento.tables:
        for linha in tabela.rows:
            linhas.extend(cell.text for cell in linha.cells)
    return "\n".join(linhas)


# --- Numeração --------------------------------------------------------------

@pytest.mark.parametrize("numero, esperado", [
    (1, "I"), (4, "IV"), (6, "VI"), (9, "IX"), (10, "X"), (14, "XIV"),
    (40, "XL"), (1987, "MCMLXXXVII"),
])
def test_romano(numero, esperado):
    assert romano(numero) == esperado


def test_romano_recusa_zero_e_negativo():
    for invalido in (0, -3):
        with pytest.raises(ValueError):
            romano(invalido)


def test_letra_dos_subitens():
    assert [letra(i) for i in range(3)] == ["a", "b", "c"]


def test_clausula_numera_itens_em_dois_niveis():
    clausula = Clausula(6, "Das hipóteses", "Observado o seguinte:", itens=[
        {"texto": "primeira faixa;"},
        {"texto": "propriedade:", "subitens": ["um;", "dois."]},
        {"texto": "terceira faixa."},
    ])
    rotulos = [rotulo for rotulo, _, _ in clausula.itens_formatados()]
    assert rotulos == ["i.", "ii.", "iii."]

    _, _, subitens = clausula.itens_formatados()[1]
    assert [rotulo for rotulo, _ in subitens] == ["a)", "b)"]


def test_cabecalho_usa_romano_maiusculo():
    clausula = Clausula(9, "Da cláusula penal", "O descumprimento...")
    assert clausula.cabecalho() == "CLÁUSULA IX – DA CLÁUSULA PENAL"


def test_item_recusa_texto_vazio():
    with pytest.raises(ValueError):
        Item(texto="   ")
    with pytest.raises(ValueError):
        Item(texto="ok", subitens=[""])


def test_itens_interpolam_os_dados_do_contrato():
    clausula = Clausula(1, "X", "Caput:", itens=[{"texto": "multa de {multa}"}])
    _, texto, _ = clausula.itens_formatados({"multa": "15%"})[0]
    assert texto == "multa de 15%"


def test_placeholder_sem_origem_nao_derruba_o_documento():
    clausula = Clausula(1, "X", "Valor de {inexistente}.")
    assert clausula.caput({"outro": "1"}) == "Valor de {inexistente}."


def test_data_por_extenso():
    assert data_por_extenso(date(2026, 8, 21)) == "21 de agosto de 2026"
    assert data_por_extenso(date(2026, 3, 1)) == "1 de março de 2026"


# --- Qualificação das partes ------------------------------------------------

def test_qualificacao_completa():
    assert _cliente_completo().qualificacao() == (
        "Isabel Terra Faillace Holanda, CPF nº 181.336.027-85, com domicílio "
        "na rua Servidão Piloto, Número 280, Campeche, Florianópolis, contato "
        "pelo telefone (48) 99656-5530 e e-mail isabelterra09@gmail.com."
    )


def test_qualificacao_de_quem_so_tem_nome_e_cpf():
    """Sem vírgula órfã, sem 'Número ,' — o trecho inteiro some."""
    assert _cliente_minimo().qualificacao() == "Só Nome, CPF nº 111.222.333-96."


@pytest.mark.parametrize("email, telefone, esperado", [
    ("a@b.com", "", ", contato pelo e-mail a@b.com."),
    ("", "(48) 9999-9999", ", contato pelo telefone (48) 9999-9999."),
    ("", "", "."),
])
def test_qualificacao_adapta_o_contato_ao_que_existe(email, telefone, esperado):
    parte = Parte("Fulano", "11122233396", "CPF", email, telefone)
    assert parte.qualificacao().endswith(esperado)


def test_endereco_parcial_nao_produz_pontuacao_orfa():
    endereco = Endereco("", "", "", "", "", "Florianópolis", "SC")
    parte = Parte("Fulano", "11122233396", "CPF", "", "", endereco)
    assert parte.qualificacao() == (
        "Fulano, CPF nº 111.222.333-96, com domicílio na Florianópolis."
    )


# --- Geração ponta a ponta --------------------------------------------------

CASOS = [
    ("cliente completo", _cliente_completo, 900.0),
    ("cliente só com nome e CPF", _cliente_minimo, 900.0),
    ("valor com centavo ímpar", _cliente_completo, 901.01),
]


@pytest.mark.parametrize("descricao, fabrica, valor", CASOS,
                         ids=[c[0] for c in CASOS])
@pytest.mark.parametrize("gerador", [PdfGenerator(), DocxGenerator()],
                         ids=["pdf", "docx"])
def test_documento_gera_sem_excecao(tmp_path, descricao, fabrica, valor, gerador):
    contrato = _contrato(fabrica(), valor=valor)
    caminho = gerador.gerar(contrato, str(tmp_path / "contrato"))
    assert open(caminho, "rb").read(4)  # arquivo não vazio


def test_pdf_tem_as_dez_clausulas_e_mais_de_uma_pagina(tmp_path):
    contrato = _contrato(_cliente_completo())
    assert len(contrato.clausulas) == 10

    caminho = PdfGenerator().gerar(contrato, str(tmp_path / "contrato"))
    conteudo = open(caminho, "rb").read()
    assert conteudo.startswith(b"%PDF")

    # `/Count N` no catálogo de páginas: o contrato não cabe numa página só,
    # então a paginação do rodapé tem o que numerar.
    paginas = int(re.search(rb"/Count (\d+)", conteudo).group(1))
    assert paginas > 1


def test_docx_traz_a_estrutura_nova(tmp_path):
    contrato = _contrato(_cliente_completo(),
                         contratado=Parte("Amanda Longhi", "98765432100",
                                          "CPF", "", "", ENDERECO_COMPLETO))
    caminho = DocxGenerator().gerar(contrato, str(tmp_path / "contrato"))
    texto = _texto_do_docx(caminho)

    assert "CLÁUSULA I – DAS DEFINIÇÕES" in texto
    assert "CLÁUSULA X – DO FORO DE ELEIÇÃO" in texto
    assert "ii. Propriedade de imagem:" in texto
    assert "a) detentor da imagem é a pessoa captada na imagem;" in texto
    assert "Florianópolis, 21 de agosto de 2026" in texto
    assert "CPF 987.654.321-00" in texto        # documento sob a assinatura
    assert "Contrato nº" not in texto           # o número saiu do documento
    assert "{" not in texto                     # nenhum placeholder cru


def test_documento_de_cliente_minimo_nao_deixa_lacuna(tmp_path):
    contrato = _contrato(_cliente_minimo())
    caminho = DocxGenerator().gerar(contrato, str(tmp_path / "contrato"))
    texto = _texto_do_docx(caminho)

    assert "com domicílio" not in texto
    assert "contato pelo" not in texto
    assert "Número ," not in texto
    assert ", ," not in texto


def test_pdf_gerado_pela_aplicacao_e_servido_do_banco(client, app):
    """O caminho completo: POST no formulário, blob no banco, download."""
    from app.models import ContractRecord
    from tests.conftest import registrar
    from tests.test_clients import DADOS_CONTRATO

    registrar(client, email="ponta-a-ponta@exemplo.com")
    client.post("/contracts/new", data=DADOS_CONTRATO, follow_redirects=True)

    with app.app_context():
        record = ContractRecord.query.one()
        assert record.pdf_data.startswith(b"%PDF")
        contrato_id = record.id

    resp = client.get(f"/contracts/{contrato_id}/download/pdf")
    assert resp.status_code == 200
    assert resp.data.startswith(b"%PDF")
    assert "contrato-isabel-terra-2026-08-21.pdf" in \
        resp.headers["Content-Disposition"]

    docx_resp = client.get(f"/contracts/{contrato_id}/download/docx")
    documento = docx.Document(io.BytesIO(docx_resp.data))
    texto = "\n".join(p.text for p in documento.paragraphs)
    assert "CLÁUSULA I – DAS DEFINIÇÕES" in texto
