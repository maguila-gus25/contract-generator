from docx import Document
from models.contract import Contrato
from .base import BaseGenerator

class DocxGenerator(BaseGenerator):
    def gerar(self, contrato: Contrato, caminho_saida: str):
        doc = Document()
        
        texto_completo = contrato.gerar_texto_completo()
        
        for linha in texto_completo.split('\n'):
            if not linha.strip():
                continue
                
            p = doc.add_paragraph()
            
            if linha == contrato.titulo.upper():
                p.add_run(linha).bold = True
                p.alignment = 1 # Center
            elif linha.startswith("CLÁUSULA") or "DO VALOR" in linha:
                p.add_run(linha).bold = True
            else:
                p.add_run(linha)
                
        doc.save(caminho_saida)
